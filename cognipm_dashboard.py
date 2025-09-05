import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))
import streamlit as st
import requests
import io
from product_agent import run_product_agent
from agent_history import run_agent_history
from roadmap_tracker import roadmap_tracker
from docx import Document
from utils.notifications import send_sms_notification
from datetime import datetime
from utils.memory import (
    load_decisions, save_decision, create_entry,
    load_prds, save_prd, clear_prds, clear_decisions,
    save_okr_analysis, load_okr_analyses, clear_okr_analyses,
    load_roadmaps, save_roadmap, clear_roadmaps
)

def query_incredible(prompt, model="small-1", chunking=True, max_tokens=1500):
    """
    Auto-continues if model ends with [CONTINUE] or hits token limit.
    Returns full completed text.
    """
    url = "https://api.incredible.one/v1/chat-completion"
    headers = {"Content-Type": "application/json"}
    messages = [
        {"role": "system", "content": (
            "You are a world-class product assistant. Always keep outputs under 1500 tokens. "
            "If more remains, end with [CONTINUE] and wait for next request."
        )},
        {"role": "user", "content": prompt}
    ]

    full_output = ""
    for _ in range(10): 
        payload = {
            "model": model,
            "max_tokens": max_tokens,
            "messages": messages
        }
        response = requests.post(url, json=payload, headers=headers)
        response.raise_for_status()
        content = response.json()["result"]["response"][0]["content"]
        full_output += content.strip().replace("[CONTINUE]", "").strip()

        if "[CONTINUE]" not in content and len(content) < (max_tokens * 4):
            break 

        messages.append({"role": "assistant", "content": content})
        messages.append({"role": "user", "content": "Continue from where you left off. Do not repeat. End with [CONTINUE] if more remains."})

    return full_output

def export_to_docx(markdown_text, filename):
    doc = Document()
    doc.add_heading(filename.replace("_", " ").title(), level=1)
    for line in markdown_text.split("\n"):
        doc.add_paragraph(line)
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

def show_dashboard():
    st.title("📄➞️🔏 From PRDs to OKRs — your intelligent assistant for every product decision")

    for key in ["generated_prd", "generated_roadmap", "optimized_roadmap", "optimized_prd", "optimized_okr"]:
        if key not in st.session_state:
            st.session_state[key] = ""

    username = st.session_state.username

    tabs = st.tabs([
        "🪳 Product Agent",
        "🛣️ Roadmap Creator",
        "📜️ Decision Logger",
        "📖 Story Intelligence",
        "🧠 Internal Coach",
        "🔍 Bias Detector",
        "🗓️ Release Readiness",
        "📐 Feature Validator",
        "📊 OKR Fit Analyzer",
        "📝 PRD Generator",
        "📄 Export Tools",
        "📚 PRD History",
        "📈 OKR Analysis History",
        "🛣️ Roadmap History",
        "🧢 Agent History",
        "🚤 Roadmap Tracker"
    ])

    with tabs[0]:
        run_product_agent()

    with tabs[1]:
        st.header("🚳️ Product Roadmap Creator")

        roadmap_vision = st.text_area("Describe your product vision or key objectives")

        def clean_roadmap_output(text: str) -> str:
            cleaned = text.strip()
            cleaned = cleaned.replace("Phase:", "### Phase")
            cleaned = cleaned.replace("Epic:", "#### Epic")
            cleaned = cleaned.replace("Features:", "**Features:**")
            cleaned = cleaned.replace("User Stories:", "**User Stories:**")
            cleaned = cleaned.replace("Acceptance Criteria:", "**Acceptance Criteria:**")

            garbage_phrases = [
                "I don't see any previous conversation",
                "Please let me know what specific work you'd like me to continue",
                "I have access to tools that can help with:"
            ]
            for phrase in garbage_phrases:
                if phrase in cleaned:
                    cleaned = cleaned.split(phrase)[0].strip()
            return cleaned

        if st.button("Generate Roadmap"):
            if roadmap_vision:
                prompt = f"""
You're a world-class AI product strategist. Based on the following **product idea**, generate a proportionate **product roadmap**.

### Instructions:
- For **small features** like "submit button", keep it short, make sure all the necessary features, user stories, and acceptance criteria is included: 1 phase, 2–3 epics max.
- All plausible phases with realistic timelines (in Q format)
- For **larger ideas**, break into multiple phases and epics.
- Assign **roles** (PM, Dev, QA, Designer, Data Scientist, etc.) for each user story
- Under each **Feature**, list:
  - **User Stories** (3–5), each starting with:  
    - _As a [role], I want to [action], so that [benefit]._  
  - **Acceptance Criteria** (3 per feature) in Gherkin style:  
    - _Given [condition], When [action], Then [result]_  
- Follow the **Markdown structure exactly as shown** below:

```markdown
### Phase 1 – Q1 2025 – [Phase Name]

#### Epic 1.1 – [Epic Name]

**Features:**
- [Feature Name]

**User Stories:**
- As a [role], I want to [action], so that [benefit].

**Acceptance Criteria:**
- Given [condition], When [action], Then [result]

⚠️ Keep your response under ~1500 tokens. If more content remains, end the response with this exact tag: `[CONTINUE]`

Product Vision:
{roadmap_vision}
"""
                raw_output = query_incredible(prompt)
                roadmap = clean_roadmap_output(raw_output)
                st.session_state["generated_roadmap"] = roadmap
                st.markdown(roadmap)

        if st.session_state["generated_roadmap"]:
            if st.button("Optimize Roadmap"):
                prompt = f"""
Improve and optimize the following product roadmap by:
- Adding more user stories per feature
- Clarifying role assignments
- Ensuring complete formatting

    Roadmap:
    {st.session_state['generated_roadmap']}
    """
                raw_output = query_incredible(prompt)
                optimized = clean_roadmap_output(raw_output)
                st.session_state["optimized_roadmap"] = optimized
                st.markdown(optimized)

            if st.button("Save Roadmap"):
                content = st.session_state["optimized_roadmap"] or st.session_state["generated_roadmap"]
                save_roadmap(content, username)
                st.success("✅ Roadmap saved!")

            if st.button("Export Roadmap as .docx"):
                content = st.session_state["optimized_roadmap"] or st.session_state["generated_roadmap"]
                doc_io = export_to_docx(content, "product_roadmap")
                st.download_button("Download Roadmap (.docx)", data=doc_io, file_name="product_roadmap.docx")


    from utils.notifications import send_sms_notification  # Make sure this is at the top of your file

    with tabs[2]:
        st.header("📜️ Decision Logger")

        # Phone input
        phone_number = st.text_input("📱 Your Phone Number (with country code)", value="+1XXXXXXXXXX")

        # Decision input
        raw_text = st.text_area("📝 Describe the decision or discussion")

        if st.button("Log Decision"):
            if raw_text:
                prompt = f"Summarize this product decision in 1–2 sentences:\n{raw_text}"
                summary = query_incredible(prompt)
                entry = create_entry(raw_text, summary, username)
                save_decision(entry)

                # Send SMS notification
                if phone_number:
                    send_sms_notification(phone_number, f"🧠 Decision Logged: {summary}")

                st.success("✅ Decision saved and SMS sent!")
                st.markdown(f"**Summary:** {summary}")
            else:
                st.warning("⚠️ Please enter a decision before submitting.")

        st.subheader("📚 Decision History")
        history = load_decisions(username)

        if history:
            for item in reversed(history):
                st.markdown(f"**🕒 {item['timestamp']}**")
                st.markdown(f"📝 {item['raw_text']}")
                st.markdown(f"🔍 _{item['summary']}_")
                st.markdown("---")
        else:
            st.info("No decisions logged yet.")

        if st.button("🗑️ Clear All Decision History"):
            clear_decisions(username)
            st.success("✅ Decision history cleared!")
            st.rerun()


    with tabs[3]:
        st.header("📖 Story Intelligence")
        prd_text = st.text_area("Paste your PRD or feature spec")
        if st.button("Analyze Spec"):
            prompt = f"""
You're a senior PM. Review this spec and provide:
- Structural improvements
- Missing info
- Suggestions to improve clarity and alignment
Spec:
{prd_text}
"""
            feedback = query_incredible(prompt)
            st.markdown("**🧠 Feedback:**")
            st.markdown(feedback)
        if st.button("Save this PRD"):
            save_prd(prd_text, username)
            st.success("✅ PRD saved to history!")

    with tabs[4]:
        st.header("🧠 Internal Product Coach")
        blocker = st.text_input("What's blocking you?")
        decision_struggle = st.text_area("Decision you're struggling with?")
        if st.button("Get Coaching"):
            prompt = f"""
You're a senior product coach. Provide advice:
Blocker: {blocker}
Struggle: {decision_struggle}
"""
            advice = query_incredible(prompt)
            st.markdown("**📌 Coaching Advice:**")
            st.markdown(advice)

    with tabs[5]:
        st.header("🔍 Bias Detector")
        bias_input = st.text_area("Paste a product spec or decision")
        if st.button("Detect Bias"):
            prompt = f"""
You are an unbiased reviewer. Identify any framing bias, confirmation bias, or emotionally loaded language in the following:
{bias_input}
"""
            result = query_incredible(prompt)
            st.markdown("**🚨 Detected Bias or Framing Issues:**")
            st.markdown(result)

    with tabs[6]:
        st.header("🗓️ Release Readiness Check")
        release_plan = st.text_area("Paste your release plan or checklist")
        if st.button("Evaluate Readiness"):
            prompt = f"""
You're a release QA advisor. Review the plan and identify gaps in:
- Test coverage
- Documentation
- Communication plan
- Rollback strategy
Plan:
{release_plan}
"""
            feedback = query_incredible(prompt)
            st.markdown("**🚪 Readiness Feedback:**")
            st.markdown(feedback)

    with tabs[7]:
        st.header("📐 Feature Validator")
        feature_idea = st.text_area("Describe the feature idea")
        if st.button("Validate Feature"):
            prompt = f"""
You are a strategic PM. Assess the following feature idea for:
- Clarity
- User alignment
- Strategic value
- Engineering feasibility
Feature:
{feature_idea}
"""
            analysis = query_incredible(prompt)
            st.markdown("**🔍 Feature Validation:**")
            st.markdown(analysis)

    with tabs[8]:
        st.header("📊 OKR Fit Analyzer")
        okrs = st.text_area("Paste your current OKRs")
        proposal = st.text_area("Describe the new feature or initiative")
        if st.button("Analyze Fit"):
            prompt = f"""
Evaluate how well this proposal aligns with the given OKRs. Be strict. Only give 9 or 10 if all key results are directly supported.
Give:
- Fit score (1–10)
- Alignment analysis
- What's missing
- Suggestions to improve fit
OKRs:
{okrs}
Proposal:
{proposal}
"""
            result = query_incredible(prompt)
            st.session_state["optimized_okr"] = result
            save_okr_analysis({
                "timestamp": str(datetime.now()),
                "username": username,
                "okrs": okrs,
                "proposal": proposal,
                "result": result
            })
            st.markdown("**📈 OKR Fit Analysis:**")
            st.markdown(result)

    with tabs[9]:
        st.header("📝 PRD Generator")
        idea = st.text_area("Describe the product idea or user problem")
        if st.button("Generate PRD"):
            prompt = f"""
Generate a structured Product Requirements Document (PRD) for the following idea. Include: Title, Problem Statement, Goals, Requirements, Metrics, Stakeholders, Timeline
Idea:
{idea}
"""
            st.session_state["generated_prd"] = query_incredible(prompt)
            st.markdown(st.session_state["generated_prd"])
        if st.session_state["generated_prd"]:
            if st.button("Optimize PRD"):
                st.session_state["optimized_prd"] = query_ollama(f"Improve and refine this PRD:\n\n{st.session_state['generated_prd']}")
                st.markdown(st.session_state["optimized_prd"])
            if st.button("Save PRD"):
                save_prd(st.session_state["generated_prd"], username)
                st.success("✅ PRD saved to history!")
            if st.button("Export PRD as .docx"):
                content = st.session_state["optimized_prd"] or st.session_state["generated_prd"]
                doc_io = export_to_docx(content, "prd_document")
                st.download_button("Download PRD (.docx)", data=doc_io, file_name="prd_document.docx")

    with tabs[10]:
        st.header("📄 Export Tools")
        export_text = st.text_area("Paste any notes, decisions, or summaries to export")
        if st.button("Export as Markdown"):
            st.download_button("Download .md", data=export_text, file_name="cognipm_export.md")
        if st.button("Export as Plain Text"):
            st.download_button("Download .txt", data=export_text, file_name="cognipm_export.txt")

    with tabs[11]:
        st.header("📚 PRD History Viewer")
        prd_history = load_prds(username)
        search_title = st.text_input("🔍 Filter by keyword in PRD title or content")
        filter_date = st.date_input("📅 Filter by date (optional)", value=None, key="filter_date")
        for item in reversed(prd_history):
            if search_title and search_title.lower() not in item["text"].lower():
                continue
            if filter_date and filter_date.strftime("%Y-%m-%d") not in item["timestamp"]:
                continue
            st.markdown(f"**🕒 {item['timestamp']}**")
            st.markdown(f"📄 {item['text']}")
            st.markdown("---")
        if st.button("🗑️ Clear All PRD History"):
            clear_prds(username)
            st.success("✅ PRD history cleared!")
            st.rerun()

    with tabs[12]:
        st.header("📈 OKR Analysis History")
        okr_analysis = load_okr_analyses(username)
        for item in reversed(okr_analysis):
            st.markdown(f"**🕒 {item['timestamp']}**")
            st.markdown(f"📋 **OKRs:** {item['okrs']}")
            st.markdown(f"📝 **Proposal:** {item['proposal']}")
            st.markdown(f"📈 **Analysis Result:** {item['result']}")
            st.markdown("---")
        if st.button("🗑️ Clear All OKR Analysis History"):
            clear_okr_analyses(username)
            st.success("✅ OKR analysis history cleared!")
            st.rerun()

    with tabs[13]:
        st.header("🛂 Roadmap History Viewer")
        roadmap_history = load_roadmaps(username)
        for item in reversed(roadmap_history):
            st.markdown(f"**🕒 {item['timestamp']}**")
            st.markdown(f"🪳 {item['text']}")
            st.markdown("---")
        if st.button("🗑️ Clear All Roadmap History"):
            clear_roadmaps(username)
            st.success("✅ Roadmap history cleared!")
            st.rerun()

    with tabs[14]:
        run_agent_history()

    with tabs[15]:
        roadmap_tracker()
